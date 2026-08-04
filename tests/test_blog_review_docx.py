import re
import struct
import zlib
import zipfile
from pathlib import Path

from docx import Document
from lxml import etree

from marin_dna.blog_review_sync import (
    PreparedFigure,
    build_review_snapshot,
    render_review_docx,
)
from marin_dna.blog_workspace import default_config_path, load_config


def _one_pixel_png() -> bytes:
    def chunk(kind: bytes, data: bytes) -> bytes:
        return (
            struct.pack(">I", len(data))
            + kind
            + data
            + struct.pack(">I", zlib.crc32(kind + data) & 0xFFFFFFFF)
        )

    header = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
    pixels = zlib.compress(b"\x00\xff\xff\xff")
    return (
        b"\x89PNG\r\n\x1a\n"
        + chunk(b"IHDR", header)
        + chunk(b"IDAT", pixels)
        + chunk(b"IEND", b"")
    )


def test_docx_has_native_structure_links_and_every_figure(tmp_path: Path) -> None:
    config = load_config(default_config_path())
    snapshot = build_review_snapshot(
        config,
        source_sha="a" * 40,
        request_id="issue-408-docx-structure",
        requested_at="2026-07-28T12:34:56Z",
    )
    prepared = []
    for figure in snapshot.figures:
        png_path = tmp_path / f"{figure.number}.png"
        png_path.write_bytes(_one_pixel_png())
        prepared.append(
            PreparedFigure(figure=figure, png_path=png_path, sha256="f" * 64)
        )

    output = tmp_path / "review.docx"
    render_review_docx(snapshot, tuple(prepared), output)
    word_namespace = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    relationship_namespace = (
        "http://schemas.openxmlformats.org/package/2006/relationships"
    )
    namespaces = {"w": word_namespace, "pr": relationship_namespace}
    with zipfile.ZipFile(output) as archive:
        names = set(archive.namelist())
        assert "word/footnotes.xml" in names
        assert "word/_rels/footnotes.xml.rels" in names
        document_root = etree.fromstring(archive.read("word/document.xml"))
        footnotes_root = etree.fromstring(archive.read("word/footnotes.xml"))
        document_relationships = etree.fromstring(
            archive.read("word/_rels/document.xml.rels")
        )
        footnote_relationships = etree.fromstring(
            archive.read("word/_rels/footnotes.xml.rels")
        )

    reference_ids = [
        int(item.get(f"{{{word_namespace}}}id"))
        for item in document_root.findall(".//w:footnoteReference", namespaces)
    ]
    native_footnotes = [
        item
        for item in footnotes_root.findall("w:footnote", namespaces)
        if int(item.get(f"{{{word_namespace}}}id")) > 0
    ]
    assert reference_ids == list(range(1, 31))
    assert [
        int(item.get(f"{{{word_namespace}}}id")) for item in native_footnotes
    ] == list(range(1, 31))
    assert all(
        item.find(".//w:footnoteRef", namespaces) is not None
        for item in native_footnotes
    )
    assert "For a broader overview" in "".join(
        native_footnotes[0].xpath(".//w:t/text()", namespaces=namespaces)
    )
    assert any(
        item.get("Type", "").endswith("/footnotes") for item in document_relationships
    )
    assert any(
        item.get("Type", "").endswith("/hyperlink") for item in footnote_relationships
    )
    assert not document_root.findall(".//w:br", namespaces)

    document = Document(output)
    headings = {
        (paragraph.style.name, paragraph.text)
        for paragraph in document.paragraphs
        if paragraph.style.name.startswith("Heading")
    }
    captions = [
        paragraph.text
        for paragraph in document.paragraphs
        if re.match(r"^Figure (?:[1-9]|1[0-9]|20):", paragraph.text)
    ]
    hyperlinks = [
        relationship
        for relationship in document.part.rels.values()
        if relationship.reltype.endswith("/hyperlink")
    ]

    assert document.paragraphs[0].style.name == "Title"
    assert document.paragraphs[0].alignment == 1
    assert ("Heading 2", "Introduction") in headings
    assert ("Heading 2", "Conclusion") in headings
    assert ("Heading 2", "Notes") not in headings
    assert len(document.inline_shapes) == 20
    assert len(captions) == 20
    assert len(hyperlinks) >= 20
