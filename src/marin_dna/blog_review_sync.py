"""Prepare immutable review documents and coordinate retry-safe blog syncs.

The canonical input remains the issue #373 Markdown workspace.  This module
turns one commit into an auditable DOCX transport bundle for Google Drive and
provides a small, dependency-injected state machine for the two external
effects required by issue #408:

* update the existing website preview branch; and
* create one new Google Doc for each distinct sync request.

Credentials and provider clients deliberately live outside this module.  An
agent can use the connected GitHub and Google Drive services, while tests use
fakes that exercise the same idempotency and partial-recovery rules.
"""

from __future__ import annotations

import argparse
import ast
import hashlib
import html
import json
import os
import re
import shutil
import tempfile
import zipfile
from dataclasses import asdict, dataclass
from datetime import datetime, timezone
from html.parser import HTMLParser
from pathlib import Path
from typing import Any, Callable, Literal, Protocol
from urllib.parse import unquote, urlsplit

from marin_dna.blog_workspace import (
    WorkspaceConfig,
    default_config_path,
    load_config,
    validate_workspace,
)


SOURCE_REPOSITORY = "https://github.com/Open-Athena/marin-dna"
SOURCE_BRANCH = "claude/issue-373-blog-staging"
WEBSITE_PREVIEW_URL = (
    "https://cms-blog-genomic-lm-optimiza.openathena-ai.pages.dev/"
    "blog/genomic-lm-optimization/"
)

FRONTMATTER_RE = re.compile(r"\A---\r?\n(?P<body>.*?)\r?\n---\r?\n", re.DOTALL)
STYLE_RE = re.compile(r"<style\b[^>]*>.*?</style>", re.IGNORECASE | re.DOTALL)
COMMENT_RE = re.compile(r"<!--.*?-->", re.DOTALL)
FIGURE_RE = re.compile(
    r"<figure\b(?P<figure_attrs>[^>]*)>\s*"
    r"<img\b(?P<img_attrs>[^>]*)/?>\s*"
    r"<figcaption>(?P<caption>.*?)</figcaption>\s*"
    r"</figure>",
    re.IGNORECASE | re.DOTALL,
)
ATTRIBUTE_RE = re.compile(
    r"""(?P<name>[A-Za-z_:][-A-Za-z0-9_:.]*)\s*=\s*
        (?:"(?P<double>[^"]*)"|'(?P<single>[^']*)')""",
    re.VERBOSE,
)
DETAILS_OPEN_RE = re.compile(
    r"<details>\s*<summary>(?P<summary>.*?)</summary>",
    re.IGNORECASE | re.DOTALL,
)
DETAILS_CLOSE_RE = re.compile(r"</details>", re.IGNORECASE)
SUPERSCRIPT_RE = re.compile(r"<sup>(?P<value>[0-9]+)</sup>", re.IGNORECASE)
FOOTNOTE_DEFINITION_RE = re.compile(
    r"^[ \t]{0,3}\[\^(?P<name>[^\]\s]+)\]:[ \t]*(?P<text>.*)$"
)
FOOTNOTE_REFERENCE_RE = re.compile(r"\[\^(?P<name>[^\]\s]+)\]")
FIGURE_MARKER_RE = re.compile(r"\A\[\[MARIN_DNA_FIGURE_(?P<number>\d{3})]]\Z")
FOOTNOTE_MARKER_RE = re.compile(r"MARINDNAFOOTNOTEREF(?P<number>\d{3})")
REQUEST_ID_RE = re.compile(r"[A-Za-z0-9][A-Za-z0-9._-]{7,127}")
SHA_RE = re.compile(r"[0-9a-f]{40}")
ISO_UTC_RE = re.compile(r"\d{4}-\d{2}-\d{2}T\d{2}:\d{2}:\d{2}Z")


@dataclass(frozen=True)
class ArticleMetadata:
    """Review-relevant values from the article frontmatter."""

    title: str
    slug: str
    author: str
    date: str
    summary: str


@dataclass(frozen=True)
class ReviewFigure:
    """One figure in canonical article order."""

    number: int
    marker: str
    source_uri: str
    source_path: Path
    alt_text: str
    caption_markdown: str


@dataclass(frozen=True)
class ReviewFootnote:
    "One native footnote in first-reference order."

    number: int
    identifier: str
    marker: str
    definition_markdown: str


@dataclass(frozen=True)
class ReviewSnapshot:
    """Provider-independent representation of one immutable review document."""

    metadata: ArticleMetadata
    source_sha: str
    request_id: str
    requested_at: str
    body_markdown: str
    figures: tuple[ReviewFigure, ...]
    footnotes: tuple[ReviewFootnote, ...]

    @property
    def source_url(self) -> str:
        """Return a commit-pinned link to the canonical article."""
        return (
            f"{SOURCE_REPOSITORY}/blob/{self.source_sha}/"
            "blog/genomic-lm-optimization/content/blog/"
            "genomic-lm-optimization.md"
        )

    @property
    def document_title(self) -> str:
        """Return a title that distinguishes immutable review snapshots."""
        return (
            f"{self.metadata.title} — review "
            f"{self.requested_at[:10]} ({self.source_sha[:12]})"
        )


@dataclass(frozen=True)
class PreparedFigure:
    """One rendered figure embedded in the DOCX transport."""

    figure: ReviewFigure
    png_path: Path
    sha256: str


@dataclass(frozen=True)
class ReviewBundle:
    """Files prepared for one Google Drive import."""

    root: Path
    markdown_path: Path
    docx_path: Path
    manifest_path: Path
    figures: tuple[PreparedFigure, ...]


@dataclass(frozen=True)
class SyncRequest:
    """Stable identity for one explicit user-triggered sync."""

    request_id: str
    source_sha: str
    requested_at: str
    bundle_sha256: str


@dataclass(frozen=True)
class WebsiteReceipt:
    """Verified result of updating the existing website preview branch."""

    commit_sha: str
    preview_url: str
    build_status: Literal["success"]


@dataclass(frozen=True)
class DocumentReceipt:
    """Verified result of creating one immutable Google Doc."""

    document_url: str
    revision_id: str
    verification_status: Literal["success"]


SyncTarget = Literal["request", "website", "document"]
SyncStatus = Literal["registered", "succeeded", "failed"]


@dataclass(frozen=True)
class SyncEvent:
    """One append-only state transition for a sync request."""

    request_id: str
    source_sha: str
    target: SyncTarget
    status: SyncStatus
    recorded_at: str
    details: dict[str, str]


@dataclass(frozen=True)
class SyncReport:
    """Current state after attempting any missing external effects."""

    request: SyncRequest
    website: WebsiteReceipt | None
    document: DocumentReceipt | None
    website_error: str | None
    document_error: str | None

    @property
    def complete(self) -> bool:
        """Return whether both required effects have verified receipts."""
        return self.website is not None and self.document is not None


class SyncLedger(Protocol):
    """Append-only storage used by the retry-safe coordinator."""

    def read(self) -> tuple[SyncEvent, ...]:
        """Return all recorded events in append order."""

    def append(self, event: SyncEvent) -> None:
        """Durably append one event."""


class WebsitePublisher(Protocol):
    """External website effect, keyed by the stable request ID."""

    def publish(self, bundle: ReviewBundle, request: SyncRequest) -> WebsiteReceipt:
        """Update or reuse the website preview and return a verified receipt."""


class DocumentPublisher(Protocol):
    """External Google Doc effect, keyed by the stable request ID."""

    def publish(self, bundle: ReviewBundle, request: SyncRequest) -> DocumentReceipt:
        """Create or recover the immutable review Doc and return its receipt."""


class _CaptionMarkdownParser(HTMLParser):
    """Convert the small HTML subset used by figure captions to Markdown."""

    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.parts: list[str] = []
        self.links: list[str] = []

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        attributes = {name: value for name, value in attrs}
        if tag == "strong":
            self.parts.append("**")
        elif tag == "em":
            self.parts.append("*")
        elif tag == "code":
            self.parts.append("`")
        elif tag == "a":
            href = attributes.get("href")
            assert href, "caption link lacks href"
            self.parts.append("[")
            self.links.append(href)
        elif tag == "br":
            self.parts.append("  \n")
        else:
            raise AssertionError(f"unsupported caption tag: {tag}")

    def handle_endtag(self, tag: str) -> None:
        if tag == "strong":
            self.parts.append("**")
        elif tag == "em":
            self.parts.append("*")
        elif tag == "code":
            self.parts.append("`")
        elif tag == "a":
            assert self.links, "caption link stack underflow"
            self.parts.append(f"]({self.links.pop()})")
        elif tag != "br":
            raise AssertionError(f"unsupported caption tag: {tag}")

    def handle_data(self, data: str) -> None:
        self.parts.append(data)

    def markdown(self) -> str:
        """Return the fully parsed caption."""
        assert not self.links, "unclosed caption link"
        return "".join(self.parts).strip()


def _parse_scalar(value: str) -> str:
    """Parse one simple YAML scalar without accepting arbitrary YAML objects."""
    stripped = value.strip()
    assert stripped, "empty frontmatter scalar"
    if stripped[0] in {'"', "'"}:
        parsed = ast.literal_eval(stripped)
        assert isinstance(parsed, str)
        return parsed
    return stripped


def split_frontmatter(markdown: str) -> tuple[ArticleMetadata, str]:
    """Extract the required metadata and return the article body."""
    match = FRONTMATTER_RE.match(markdown)
    assert match, "article must begin with YAML frontmatter"
    fields: dict[str, str] = {}
    for line in match.group("body").splitlines():
        if line.startswith((" ", "\t", "- ")):
            continue
        key, separator, value = line.partition(":")
        assert separator, f"invalid frontmatter line: {line!r}"
        if key in {"title", "slug", "author", "date", "summary"}:
            assert key not in fields, f"duplicate frontmatter field: {key}"
            fields[key] = _parse_scalar(value)
    required = {"title", "slug", "author", "date", "summary"}
    assert set(fields) == required, (
        f"review frontmatter fields differ: {set(fields) ^ required}"
    )
    datetime.strptime(fields["date"], "%Y-%m-%d")
    metadata = ArticleMetadata(
        title=fields["title"],
        slug=fields["slug"],
        author=fields["author"],
        date=fields["date"],
        summary=fields["summary"],
    )
    return metadata, markdown[match.end() :]


def _parse_attributes(raw: str) -> dict[str, str]:
    """Parse quoted attributes and reject malformed leftovers."""
    attributes: dict[str, str] = {}
    spans: list[tuple[int, int]] = []
    for match in ATTRIBUTE_RE.finditer(raw):
        name = match.group("name").lower()
        assert name not in attributes, f"duplicate HTML attribute: {name}"
        attributes[name] = html.unescape(
            match.group("double")
            if match.group("double") is not None
            else match.group("single")
        )
        spans.append(match.span())
    remainder = raw
    for start, end in reversed(spans):
        remainder = remainder[:start] + remainder[end:]
    assert not remainder.strip().rstrip("/").strip(), (
        f"unsupported HTML attributes: {remainder!r}"
    )
    return attributes


def _caption_to_markdown(raw: str) -> str:
    parser = _CaptionMarkdownParser()
    parser.feed(raw)
    parser.close()
    return parser.markdown()


def _resolve_figure_path(config: WorkspaceConfig, source_uri: str) -> Path:
    """Resolve a figure URL and assert that it stays in the canonical asset tree."""
    parsed = urlsplit(source_uri)
    assert not parsed.scheme and not parsed.netloc, source_uri
    decoded = unquote(parsed.path)
    expected_prefix = f"/assets/images/blog/{config.slug}/"
    assert decoded.startswith(expected_prefix), (
        f"review figure is outside canonical assets: {source_uri}"
    )
    relative = Path(decoded.removeprefix(expected_prefix))
    assert relative.parts and ".." not in relative.parts, source_uri
    source = (config.assets / relative).resolve()
    assert source.is_relative_to(config.assets.resolve()), source
    assert source.is_file(), f"missing review figure: {source}"
    return source


def _extract_figures(
    body: str, config: WorkspaceConfig
) -> tuple[str, tuple[ReviewFigure, ...]]:
    figures: list[ReviewFigure] = []

    def replace(match: re.Match[str]) -> str:
        figure_attributes = _parse_attributes(match.group("figure_attrs"))
        assert set(figure_attributes) == {"id"}, (
            f"figure attributes differ: {set(figure_attributes)}"
        )
        image_attributes = _parse_attributes(match.group("img_attrs"))
        assert set(image_attributes) == {"src", "alt"}, (
            f"image attributes differ: {set(image_attributes)}"
        )
        number = len(figures) + 1
        marker = f"[[MARIN_DNA_FIGURE_{number:03d}]]"
        caption = _caption_to_markdown(match.group("caption"))
        assert caption.startswith(f"**Figure {number}:**"), (
            f"figure caption numbering differs at figure {number}: {caption}"
        )
        figures.append(
            ReviewFigure(
                number=number,
                marker=marker,
                source_uri=image_attributes["src"],
                source_path=_resolve_figure_path(config, image_attributes["src"]),
                alt_text=image_attributes["alt"],
                caption_markdown=caption,
            )
        )
        return f"\n\n{marker}\n\n"

    replaced = FIGURE_RE.sub(replace, body)
    assert "<figure" not in replaced.lower(), "unparsed figure HTML remains"
    assert figures, "article has no review figures"
    return replaced, tuple(figures)


def _extract_footnotes(markdown: str) -> tuple[str, dict[str, str]]:
    """Remove Markdown footnote definitions, including indented continuations."""
    lines = markdown.splitlines()
    body: list[str] = []
    definitions: dict[str, str] = {}
    index = 0
    while index < len(lines):
        match = FOOTNOTE_DEFINITION_RE.match(lines[index])
        if match is None:
            body.append(lines[index])
            index += 1
            continue

        name = match.group("name")
        assert name not in definitions, f"duplicate footnote identifier: {name}"
        content = [match.group("text")]
        index += 1
        while index < len(lines):
            line = lines[index]
            if line.startswith("    "):
                content.append(line[4:])
                index += 1
                continue
            if line.startswith("\t"):
                content.append(line[1:])
                index += 1
                continue
            if (
                not line
                and index + 1 < len(lines)
                and lines[index + 1].startswith(("    ", "\t"))
            ):
                content.append("")
                index += 1
                continue
            break
        definitions[name] = "\n".join(content).strip()
    return "\n".join(body), definitions


def _materialize_footnotes(
    markdown: str,
) -> tuple[str, tuple[ReviewFootnote, ...]]:
    "Replace Markdown references with stable native-footnote markers."
    without_definitions, definitions = _extract_footnotes(markdown)
    ordered_names: list[str] = []
    for match in FOOTNOTE_REFERENCE_RE.finditer(without_definitions):
        name = match.group("name")
        assert name in definitions, f"undefined footnote identifier: {name}"
        if name not in ordered_names:
            ordered_names.append(name)
    unreferenced = set(definitions) - set(ordered_names)
    assert not unreferenced, f"unreferenced footnotes: {sorted(unreferenced)}"
    footnotes = tuple(
        ReviewFootnote(
            number=index,
            identifier=name,
            marker=f"MARINDNAFOOTNOTEREF{index:03d}",
            definition_markdown=re.sub(r"\s*\n\s*", " ", definitions[name]).strip(),
        )
        for index, name in enumerate(ordered_names, start=1)
    )
    markers = {footnote.identifier: footnote.marker for footnote in footnotes}
    body = FOOTNOTE_REFERENCE_RE.sub(
        lambda match: markers[match.group("name")], without_definitions
    )
    assert not FOOTNOTE_REFERENCE_RE.search(body)
    return body, footnotes


def _replace_details(markdown: str) -> str:
    """Make collapsed website-only details visible in a linear review Doc."""

    def replace(match: re.Match[str]) -> str:
        summary = re.sub(r"\s+", " ", html.unescape(match.group("summary"))).strip()
        assert summary
        return f"\n\n### Supplemental figure: {summary}\n\n"

    replaced = DETAILS_OPEN_RE.sub(replace, markdown)
    replaced = DETAILS_CLOSE_RE.sub("", replaced)
    assert "<details" not in replaced.lower()
    assert "<summary" not in replaced.lower()
    return replaced


def _replace_superscripts(markdown: str) -> str:
    """Replace the article's numeric HTML superscripts with Unicode text."""
    translation = str.maketrans("0123456789", "⁰¹²³⁴⁵⁶⁷⁸⁹")
    return SUPERSCRIPT_RE.sub(
        lambda match: match.group("value").translate(translation), markdown
    )


def build_review_snapshot(
    config: WorkspaceConfig,
    *,
    source_sha: str,
    request_id: str,
    requested_at: str,
) -> ReviewSnapshot:
    """Parse and normalize the canonical article for one review request."""
    assert SHA_RE.fullmatch(source_sha), f"invalid source SHA: {source_sha}"
    assert REQUEST_ID_RE.fullmatch(request_id), f"invalid request ID: {request_id}"
    assert ISO_UTC_RE.fullmatch(requested_at), (
        f"requested_at must be second-precision UTC: {requested_at}"
    )
    datetime.strptime(requested_at, "%Y-%m-%dT%H:%M:%SZ")
    validate_workspace(config)
    metadata, body = split_frontmatter(config.article.read_text())
    assert metadata.slug == config.slug

    styles = STYLE_RE.findall(body)
    assert len(styles) == 1, f"expected one article style block, found {len(styles)}"
    body = STYLE_RE.sub("", body)
    body = COMMENT_RE.sub("", body)
    body, figures = _extract_figures(body, config)
    body = _replace_details(body)
    body = _replace_superscripts(body)
    body, footnotes = _materialize_footnotes(body)
    body = re.sub(r"\n{3,}", "\n\n", body).strip() + "\n"
    assert not re.search(r"<(?:style|figure|img|figcaption|details|summary)\b", body)
    assert tuple(figure.number for figure in figures) == tuple(
        range(1, len(figures) + 1)
    )
    return ReviewSnapshot(
        metadata=metadata,
        source_sha=source_sha,
        request_id=request_id,
        requested_at=requested_at,
        body_markdown=body,
        figures=figures,
        footnotes=footnotes,
    )


def snapshot_markdown(snapshot: ReviewSnapshot) -> str:
    """Render a human-inspectable Markdown companion to the DOCX transport."""
    provenance = (
        f"> **Immutable review snapshot.** Source: "
        f"[`Open-Athena/marin-dna@{snapshot.source_sha[:12]}`]"
        f"({snapshot.source_url}); request `{snapshot.request_id}`; generated "
        f"{snapshot.requested_at}. The staging Markdown on `{SOURCE_BRANCH}` is "
        "authoritative. Comments and edits in this Google Doc are review input "
        "and are not written back automatically."
    )
    body = snapshot.body_markdown
    for footnote in snapshot.footnotes:
        body = body.replace(footnote.marker, f"[^{footnote.identifier}]")
    definitions = "\n\n".join(
        f"[^{footnote.identifier}]: {footnote.definition_markdown}"
        for footnote in snapshot.footnotes
    )
    if definitions:
        body = f"{body.rstrip()}\n\n{definitions}\n"
    return (
        f"# {snapshot.metadata.title}\n\n"
        f"**{snapshot.metadata.author}**\n\n"
        f"{snapshot.metadata.date}\n\n"
        f"*{snapshot.metadata.summary}*\n\n"
        f"{provenance}\n\n"
        f"{body}"
    )


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def render_figure_png(figure: ReviewFigure, destination: Path) -> None:
    """Render one SVG figure to a high-resolution PNG for Google Docs."""
    assert figure.source_path.suffix.lower() == ".svg", figure.source_path
    from cairosvg import svg2png

    destination.parent.mkdir(parents=True, exist_ok=True)
    svg2png(
        url=str(figure.source_path),
        write_to=str(destination),
        output_width=1800,
    )
    assert destination.read_bytes().startswith(b"\x89PNG\r\n\x1a\n"), destination


def _add_hyperlink(paragraph: Any, text: str, url: str) -> None:
    """Append an external hyperlink to a python-docx paragraph."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1155CC")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.extend((color, underline))
    run.append(properties)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _inline_plain_text(nodes: list[dict[str, Any]]) -> str:
    """Return inline token text for figure-marker detection."""
    parts: list[str] = []
    for node in nodes:
        node_type = node["type"]
        if node_type in {"text", "codespan"}:
            parts.append(node["raw"])
        elif node_type == "softbreak":
            parts.append("\n")
        elif "children" in node:
            parts.append(_inline_plain_text(node["children"]))
        else:
            return ""
    return "".join(parts)


def _add_footnote_reference(paragraph: Any, number: int) -> None:
    "Append one native Word footnote reference to a paragraph."
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    style = OxmlElement("w:rStyle")
    style.set(qn("w:val"), "FootnoteReference")
    properties.append(style)
    run.append(properties)
    reference = OxmlElement("w:footnoteReference")
    reference.set(qn("w:id"), str(number))
    run.append(reference)
    paragraph._p.append(run)


def _render_inline(
    paragraph: Any,
    nodes: list[dict[str, Any]],
    footnotes: dict[int, ReviewFootnote],
    *,
    bold: bool = False,
    italic: bool = False,
) -> None:
    "Render supported Mistune inline tokens into one DOCX paragraph."
    for node in nodes:
        node_type = node["type"]
        if node_type == "text":
            raw = node["raw"]
            cursor = 0
            for match in FOOTNOTE_MARKER_RE.finditer(raw):
                if match.start() > cursor:
                    run = paragraph.add_run(raw[cursor : match.start()])
                    run.bold = bold
                    run.italic = italic
                number = int(match.group("number"))
                assert number in footnotes, f"unknown footnote marker: {number}"
                _add_footnote_reference(paragraph, number)
                cursor = match.end()
            if cursor < len(raw):
                run = paragraph.add_run(raw[cursor:])
                run.bold = bold
                run.italic = italic
        elif node_type == "softbreak":
            paragraph.add_run(" ")
        elif node_type == "linebreak":
            paragraph.add_run().add_break()
        elif node_type == "strong":
            _render_inline(
                paragraph,
                node["children"],
                footnotes,
                bold=True,
                italic=italic,
            )
        elif node_type == "emphasis":
            _render_inline(
                paragraph,
                node["children"],
                footnotes,
                bold=bold,
                italic=True,
            )
        elif node_type == "codespan":
            run = paragraph.add_run(node["raw"])
            run.bold = bold
            run.italic = italic
            run.font.name = "Aptos Mono"
        elif node_type == "link":
            label = _inline_plain_text(node["children"])
            assert label
            assert not FOOTNOTE_MARKER_RE.search(label)
            _add_hyperlink(paragraph, label, node["attrs"]["url"])
        else:
            raise AssertionError(f"unsupported inline Markdown token: {node_type}")


def _render_blocks(
    document: Any,
    nodes: list[dict[str, Any]],
    figures: dict[int, PreparedFigure],
    footnotes: dict[int, ReviewFootnote],
    *,
    list_style: str | None = None,
) -> None:
    """Render supported Mistune block tokens into a native DOCX structure."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches

    for node in nodes:
        node_type = node["type"]
        if node_type == "blank_line":
            continue
        if node_type in {"paragraph", "block_text"}:
            children = node["children"]
            marker_match = FIGURE_MARKER_RE.fullmatch(_inline_plain_text(children))
            if marker_match:
                number = int(marker_match.group("number"))
                prepared = figures[number]
                image_paragraph = document.add_paragraph()
                image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = image_paragraph.add_run()
                run.add_picture(str(prepared.png_path), width=Inches(6.25))
                caption = document.add_paragraph()
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption_nodes = _parse_markdown_inlines(
                    prepared.figure.caption_markdown
                )
                _render_inline(caption, caption_nodes, footnotes)
                continue
            paragraph = document.add_paragraph(style=list_style)
            _render_inline(paragraph, children, footnotes)
        elif node_type == "heading":
            level = node["attrs"]["level"]
            assert 1 <= level <= 3, f"unsupported heading level: {level}"
            paragraph = document.add_heading(level=level)
            _render_inline(paragraph, node["children"], footnotes)
        elif node_type == "block_quote":
            before = len(document.paragraphs)
            _render_blocks(document, node["children"], figures, footnotes)
            for paragraph in document.paragraphs[before:]:
                paragraph.style = "Quote"
        elif node_type == "list":
            depth = node["attrs"]["depth"]
            assert depth <= 2, f"unsupported list depth: {depth}"
            base = "List Number" if node["attrs"]["ordered"] else "List Bullet"
            style = base if depth == 0 else f"{base} {depth + 1}"
            for item in node["children"]:
                assert item["type"] == "list_item"
                first = True
                for child in item["children"]:
                    if child["type"] in {"block_text", "paragraph"}:
                        _render_blocks(
                            document,
                            [child],
                            figures,
                            footnotes,
                            list_style=style if first else None,
                        )
                        first = False
                    elif child["type"] == "list":
                        _render_blocks(document, [child], figures, footnotes)
                    else:
                        raise AssertionError(f"unsupported list child: {child['type']}")
        elif node_type == "block_code":
            paragraph = document.add_paragraph()
            run = paragraph.add_run(node["raw"].rstrip())
            run.font.name = "Aptos Mono"
        elif node_type == "thematic_break":
            paragraph = document.add_paragraph()
            paragraph.add_run("—")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            raise AssertionError(f"unsupported block Markdown token: {node_type}")


def _parse_markdown(markdown: str) -> list[dict[str, Any]]:
    import mistune

    parser = mistune.create_markdown(renderer=None)
    parsed = parser(markdown)
    assert isinstance(parsed, list)
    return parsed


def _parse_markdown_inlines(markdown: str) -> list[dict[str, Any]]:
    parsed = _parse_markdown(markdown)
    blocks = [node for node in parsed if node["type"] != "blank_line"]
    assert len(blocks) == 1 and blocks[0]["type"] == "paragraph", blocks
    return blocks[0]["children"]


_WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_OFFICE_REL_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
_PACKAGE_REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
_CONTENT_TYPE_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
_XML_NS = "http://www.w3.org/XML/1998/namespace"


def _xml_name(namespace: str, local: str) -> str:
    return f"{{{namespace}}}{local}"


def _append_ooxml_run(
    parent: Any,
    text: str,
    *,
    bold: bool = False,
    italic: bool = False,
    monospace: bool = False,
    link_style: bool = False,
) -> None:
    from lxml import etree

    run = etree.SubElement(parent, _xml_name(_WORD_NS, "r"))
    if bold or italic or monospace or link_style:
        properties = etree.SubElement(run, _xml_name(_WORD_NS, "rPr"))
        if bold:
            etree.SubElement(properties, _xml_name(_WORD_NS, "b"))
        if italic:
            etree.SubElement(properties, _xml_name(_WORD_NS, "i"))
        if monospace:
            fonts = etree.SubElement(properties, _xml_name(_WORD_NS, "rFonts"))
            fonts.set(_xml_name(_WORD_NS, "ascii"), "Aptos Mono")
            fonts.set(_xml_name(_WORD_NS, "hAnsi"), "Aptos Mono")
        if link_style:
            color = etree.SubElement(properties, _xml_name(_WORD_NS, "color"))
            color.set(_xml_name(_WORD_NS, "val"), "1155CC")
            underline = etree.SubElement(properties, _xml_name(_WORD_NS, "u"))
            underline.set(_xml_name(_WORD_NS, "val"), "single")
    text_element = etree.SubElement(run, _xml_name(_WORD_NS, "t"))
    text_element.set(_xml_name(_XML_NS, "space"), "preserve")
    text_element.text = text


def _append_footnote_inlines(
    parent: Any,
    nodes: list[dict[str, Any]],
    relationships: list[tuple[str, str]],
    *,
    bold: bool = False,
    italic: bool = False,
    link_style: bool = False,
) -> None:
    from lxml import etree

    for node in nodes:
        node_type = node["type"]
        if node_type == "text":
            assert not FOOTNOTE_MARKER_RE.search(node["raw"])
            _append_ooxml_run(
                parent, node["raw"], bold=bold, italic=italic, link_style=link_style
            )
        elif node_type == "softbreak":
            _append_ooxml_run(parent, " ", link_style=link_style)
        elif node_type == "linebreak":
            run = etree.SubElement(parent, _xml_name(_WORD_NS, "r"))
            etree.SubElement(run, _xml_name(_WORD_NS, "br"))
        elif node_type == "strong":
            _append_footnote_inlines(
                parent,
                node["children"],
                relationships,
                bold=True,
                italic=italic,
                link_style=link_style,
            )
        elif node_type == "emphasis":
            _append_footnote_inlines(
                parent,
                node["children"],
                relationships,
                bold=bold,
                italic=True,
                link_style=link_style,
            )
        elif node_type == "codespan":
            _append_ooxml_run(
                parent,
                node["raw"],
                bold=bold,
                italic=italic,
                monospace=True,
                link_style=link_style,
            )
        elif node_type == "link":
            relationship_id = f"rId{len(relationships) + 1}"
            relationships.append((relationship_id, node["attrs"]["url"]))
            hyperlink = etree.SubElement(parent, _xml_name(_WORD_NS, "hyperlink"))
            hyperlink.set(_xml_name(_OFFICE_REL_NS, "id"), relationship_id)
            _append_footnote_inlines(
                hyperlink,
                node["children"],
                relationships,
                bold=bold,
                italic=italic,
                link_style=True,
            )
        else:
            raise AssertionError(f"unsupported footnote Markdown token: {node_type}")


def _footnotes_xml(
    footnotes: tuple[ReviewFootnote, ...],
) -> tuple[bytes, list[tuple[str, str]]]:
    from lxml import etree

    root = etree.Element(
        _xml_name(_WORD_NS, "footnotes"), nsmap={"w": _WORD_NS, "r": _OFFICE_REL_NS}
    )
    for footnote_id, footnote_type, separator_name in (
        (-1, "separator", "separator"),
        (0, "continuationSeparator", "continuationSeparator"),
    ):
        footnote = etree.SubElement(root, _xml_name(_WORD_NS, "footnote"))
        footnote.set(_xml_name(_WORD_NS, "id"), str(footnote_id))
        footnote.set(_xml_name(_WORD_NS, "type"), footnote_type)
        paragraph = etree.SubElement(footnote, _xml_name(_WORD_NS, "p"))
        run = etree.SubElement(paragraph, _xml_name(_WORD_NS, "r"))
        etree.SubElement(run, _xml_name(_WORD_NS, separator_name))

    relationships: list[tuple[str, str]] = []
    for expected_number, item in enumerate(footnotes, start=1):
        assert item.number == expected_number
        footnote = etree.SubElement(root, _xml_name(_WORD_NS, "footnote"))
        footnote.set(_xml_name(_WORD_NS, "id"), str(item.number))
        paragraph = etree.SubElement(footnote, _xml_name(_WORD_NS, "p"))
        paragraph_properties = etree.SubElement(paragraph, _xml_name(_WORD_NS, "pPr"))
        paragraph_style = etree.SubElement(
            paragraph_properties, _xml_name(_WORD_NS, "pStyle")
        )
        paragraph_style.set(_xml_name(_WORD_NS, "val"), "FootnoteText")
        reference_run = etree.SubElement(paragraph, _xml_name(_WORD_NS, "r"))
        reference_properties = etree.SubElement(
            reference_run, _xml_name(_WORD_NS, "rPr")
        )
        reference_style = etree.SubElement(
            reference_properties, _xml_name(_WORD_NS, "rStyle")
        )
        reference_style.set(_xml_name(_WORD_NS, "val"), "FootnoteReference")
        etree.SubElement(reference_run, _xml_name(_WORD_NS, "footnoteRef"))
        _append_ooxml_run(paragraph, " ")
        _append_footnote_inlines(
            paragraph, _parse_markdown_inlines(item.definition_markdown), relationships
        )
    return etree.tostring(
        root, xml_declaration=True, encoding="UTF-8", standalone=True
    ), relationships


def _add_footnote_styles(styles_xml: bytes) -> bytes:
    from lxml import etree

    root = etree.fromstring(styles_xml)
    existing = {
        style.get(_xml_name(_WORD_NS, "styleId"))
        for style in root.findall(_xml_name(_WORD_NS, "style"))
    }
    if "FootnoteText" not in existing:
        style = etree.SubElement(root, _xml_name(_WORD_NS, "style"))
        style.set(_xml_name(_WORD_NS, "type"), "paragraph")
        style.set(_xml_name(_WORD_NS, "styleId"), "FootnoteText")
        name = etree.SubElement(style, _xml_name(_WORD_NS, "name"))
        name.set(_xml_name(_WORD_NS, "val"), "footnote text")
        based_on = etree.SubElement(style, _xml_name(_WORD_NS, "basedOn"))
        based_on.set(_xml_name(_WORD_NS, "val"), "Normal")
        run_properties = etree.SubElement(style, _xml_name(_WORD_NS, "rPr"))
        size = etree.SubElement(run_properties, _xml_name(_WORD_NS, "sz"))
        size.set(_xml_name(_WORD_NS, "val"), "20")
        size_complex = etree.SubElement(run_properties, _xml_name(_WORD_NS, "szCs"))
        size_complex.set(_xml_name(_WORD_NS, "val"), "20")
    if "FootnoteReference" not in existing:
        style = etree.SubElement(root, _xml_name(_WORD_NS, "style"))
        style.set(_xml_name(_WORD_NS, "type"), "character")
        style.set(_xml_name(_WORD_NS, "styleId"), "FootnoteReference")
        name = etree.SubElement(style, _xml_name(_WORD_NS, "name"))
        name.set(_xml_name(_WORD_NS, "val"), "footnote reference")
        based_on = etree.SubElement(style, _xml_name(_WORD_NS, "basedOn"))
        based_on.set(_xml_name(_WORD_NS, "val"), "DefaultParagraphFont")
        run_properties = etree.SubElement(style, _xml_name(_WORD_NS, "rPr"))
        vertical = etree.SubElement(run_properties, _xml_name(_WORD_NS, "vertAlign"))
        vertical.set(_xml_name(_WORD_NS, "val"), "superscript")
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)


def _add_native_footnotes_to_docx(
    destination: Path,
    footnotes: tuple[ReviewFootnote, ...],
) -> None:
    "Add the OOXML parts that make imported references native footnotes."
    if not footnotes:
        return
    from lxml import etree

    with zipfile.ZipFile(destination) as archive:
        entries = [(item, archive.read(item.filename)) for item in archive.infolist()]
    by_name = {item.filename: data for item, data in entries}
    required = {
        "[Content_Types].xml",
        "word/_rels/document.xml.rels",
        "word/styles.xml",
    }
    assert required <= set(by_name)

    content_types = etree.fromstring(by_name["[Content_Types].xml"])
    overrides = content_types.findall(_xml_name(_CONTENT_TYPE_NS, "Override"))
    assert not any(item.get("PartName") == "/word/footnotes.xml" for item in overrides)
    override = etree.SubElement(content_types, _xml_name(_CONTENT_TYPE_NS, "Override"))
    override.set("PartName", "/word/footnotes.xml")
    override.set(
        "ContentType",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml",
    )

    document_relationships = etree.fromstring(by_name["word/_rels/document.xml.rels"])
    footnote_relationship_type = f"{_OFFICE_REL_NS}/footnotes"
    assert not any(
        item.get("Type") == footnote_relationship_type
        for item in document_relationships
    )
    used_ids = {item.get("Id") for item in document_relationships}
    relationship_id = "rIdFootnotes"
    assert relationship_id not in used_ids
    relationship = etree.SubElement(
        document_relationships, _xml_name(_PACKAGE_REL_NS, "Relationship")
    )
    relationship.set("Id", relationship_id)
    relationship.set("Type", footnote_relationship_type)
    relationship.set("Target", "footnotes.xml")

    footnotes_xml, hyperlink_relationships = _footnotes_xml(footnotes)
    footnote_rels = etree.Element(
        _xml_name(_PACKAGE_REL_NS, "Relationships"), nsmap={None: _PACKAGE_REL_NS}
    )
    for relationship_id, target in hyperlink_relationships:
        relationship = etree.SubElement(
            footnote_rels, _xml_name(_PACKAGE_REL_NS, "Relationship")
        )
        relationship.set("Id", relationship_id)
        relationship.set("Type", f"{_OFFICE_REL_NS}/hyperlink")
        relationship.set("Target", target)
        relationship.set("TargetMode", "External")

    modified = {
        "[Content_Types].xml": etree.tostring(
            content_types, xml_declaration=True, encoding="UTF-8", standalone=True
        ),
        "word/_rels/document.xml.rels": etree.tostring(
            document_relationships,
            xml_declaration=True,
            encoding="UTF-8",
            standalone=True,
        ),
        "word/styles.xml": _add_footnote_styles(by_name["word/styles.xml"]),
    }
    file_descriptor, temporary_name = tempfile.mkstemp(
        prefix=f".{destination.name}.", suffix=".tmp", dir=destination.parent
    )
    os.close(file_descriptor)
    temporary = Path(temporary_name)
    try:
        with zipfile.ZipFile(temporary, "w") as output:
            for item, data in entries:
                output.writestr(item, modified.get(item.filename, data))
            output.writestr("word/footnotes.xml", footnotes_xml, zipfile.ZIP_DEFLATED)
            if hyperlink_relationships:
                output.writestr(
                    "word/_rels/footnotes.xml.rels",
                    etree.tostring(
                        footnote_rels,
                        xml_declaration=True,
                        encoding="UTF-8",
                        standalone=True,
                    ),
                    zipfile.ZIP_DEFLATED,
                )
        os.replace(temporary, destination)
    finally:
        temporary.unlink(missing_ok=True)


def render_review_docx(
    snapshot: ReviewSnapshot,
    figures: tuple[PreparedFigure, ...],
    destination: Path,
) -> None:
    """Render the normalized snapshot as a self-contained DOCX transport."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt

    assert tuple(item.figure for item in figures) == snapshot.figures
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(12)

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run(snapshot.metadata.title)
    for value, italic in (
        (snapshot.metadata.author, False),
        (snapshot.metadata.date, False),
        (snapshot.metadata.summary, True),
    ):
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(value)
        run.italic = italic

    table = document.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "E8F0FE")
    cell._tc.get_or_add_tcPr().append(shading)
    provenance = cell.paragraphs[0]
    provenance.add_run("Immutable review snapshot. ").bold = True
    provenance.add_run("Authoritative source: ")
    _add_hyperlink(
        provenance,
        f"Open-Athena/marin-dna@{snapshot.source_sha[:12]}",
        snapshot.source_url,
    )
    provenance.add_run(
        f"; request {snapshot.request_id}; generated {snapshot.requested_at}. "
        f"The Markdown on {SOURCE_BRANCH} is authoritative. Edits in this Doc "
        "are review input and are not written back automatically."
    )

    parsed = _parse_markdown(snapshot.body_markdown)
    figure_map = {item.figure.number: item for item in figures}
    assert set(figure_map) == set(range(1, len(figures) + 1))
    footnote_map = {item.number: item for item in snapshot.footnotes}
    assert set(footnote_map) == set(range(1, len(footnote_map) + 1))
    _render_blocks(document, parsed, figure_map, footnote_map)
    destination.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(destination))
    _add_native_footnotes_to_docx(destination, snapshot.footnotes)
    assert destination.read_bytes().startswith(b"PK\x03\x04"), destination


def _bundle_manifest(
    snapshot: ReviewSnapshot,
    *,
    markdown_path: Path,
    docx_path: Path,
    figures: tuple[PreparedFigure, ...],
    root: Path,
) -> dict[str, Any]:
    return {
        "schema_version": 1,
        "request_id": snapshot.request_id,
        "source_sha": snapshot.source_sha,
        "source_url": snapshot.source_url,
        "requested_at": snapshot.requested_at,
        "document_title": snapshot.document_title,
        "footnotes": [
            {
                "number": item.number,
                "identifier": item.identifier,
                "definition_markdown": item.definition_markdown,
            }
            for item in snapshot.footnotes
        ],
        "files": {
            markdown_path.relative_to(root).as_posix(): _sha256(markdown_path),
            docx_path.relative_to(root).as_posix(): _sha256(docx_path),
            **{
                item.png_path.relative_to(root).as_posix(): item.sha256
                for item in figures
            },
        },
        "figures": [
            {
                "number": item.figure.number,
                "marker": item.figure.marker,
                "source_uri": item.figure.source_uri,
                "source_sha256": _sha256(item.figure.source_path),
                "png": item.png_path.relative_to(root).as_posix(),
                "png_sha256": item.sha256,
                "alt_text": item.figure.alt_text,
                "caption_markdown": item.figure.caption_markdown,
            }
            for item in figures
        ],
    }


def _load_existing_bundle(root: Path) -> ReviewBundle:
    manifest_path = root / "manifest.json"
    assert manifest_path.is_file(), f"incomplete existing review bundle: {root}"
    manifest = json.loads(manifest_path.read_text())
    assert manifest["schema_version"] == 1
    for relative, expected_sha in manifest["files"].items():
        path = root / relative
        assert path.is_file(), f"review bundle file missing: {path}"
        assert _sha256(path) == expected_sha, f"review bundle file changed: {path}"
    figures = tuple(
        PreparedFigure(
            figure=ReviewFigure(
                number=raw["number"],
                marker=raw["marker"],
                source_uri=raw["source_uri"],
                source_path=Path(),
                alt_text=raw["alt_text"],
                caption_markdown=raw["caption_markdown"],
            ),
            png_path=root / raw["png"],
            sha256=raw["png_sha256"],
        )
        for raw in manifest["figures"]
    )
    return ReviewBundle(
        root=root,
        markdown_path=root / "review.md",
        docx_path=root / "review.docx",
        manifest_path=manifest_path,
        figures=figures,
    )


def prepare_review_bundle(
    config: WorkspaceConfig,
    *,
    source_sha: str,
    request_id: str,
    requested_at: str,
    destination: Path,
    figure_renderer: Callable[[ReviewFigure, Path], None] = render_figure_png,
    docx_renderer: Callable[
        [ReviewSnapshot, tuple[PreparedFigure, ...], Path], None
    ] = render_review_docx,
) -> ReviewBundle:
    """Atomically prepare or verify the immutable bundle for one request."""
    destination = destination.resolve()
    if destination.exists():
        existing = _load_existing_bundle(destination)
        manifest = json.loads(existing.manifest_path.read_text())
        assert manifest["request_id"] == request_id, "request ID bundle conflict"
        assert manifest["source_sha"] == source_sha, "source SHA bundle conflict"
        assert manifest["requested_at"] == requested_at, "request time bundle conflict"
        return existing

    snapshot = build_review_snapshot(
        config,
        source_sha=source_sha,
        request_id=request_id,
        requested_at=requested_at,
    )
    destination.parent.mkdir(parents=True, exist_ok=True)
    temporary = Path(
        tempfile.mkdtemp(
            prefix=f".{destination.name}.next-",
            dir=destination.parent,
        )
    )
    try:
        markdown_path = temporary / "review.md"
        markdown_path.write_text(snapshot_markdown(snapshot))
        prepared_figures: list[PreparedFigure] = []
        for figure in snapshot.figures:
            png_path = temporary / "figures" / f"figure-{figure.number:03d}.png"
            figure_renderer(figure, png_path)
            assert png_path.is_file(), f"figure renderer did not create {png_path}"
            prepared_figures.append(
                PreparedFigure(
                    figure=figure,
                    png_path=png_path,
                    sha256=_sha256(png_path),
                )
            )
        figures = tuple(prepared_figures)
        docx_path = temporary / "review.docx"
        docx_renderer(snapshot, figures, docx_path)
        assert docx_path.is_file(), f"DOCX renderer did not create {docx_path}"
        manifest = _bundle_manifest(
            snapshot,
            markdown_path=markdown_path,
            docx_path=docx_path,
            figures=figures,
            root=temporary,
        )
        manifest_path = temporary / "manifest.json"
        manifest_path.write_text(json.dumps(manifest, indent=2, sort_keys=True) + "\n")
        os.replace(temporary, destination)
    except Exception:
        shutil.rmtree(temporary, ignore_errors=True)
        raise
    return _load_existing_bundle(destination)


class JsonlSyncLedger:
    """Local append-only JSONL implementation used by agents and tests."""

    def __init__(self, path: Path) -> None:
        self.path = path

    def read(self) -> tuple[SyncEvent, ...]:
        """Read and validate every complete event."""
        if not self.path.exists():
            return ()
        events: list[SyncEvent] = []
        for line_number, line in enumerate(self.path.read_text().splitlines(), start=1):
            raw = json.loads(line)
            assert set(raw) == {
                "request_id",
                "source_sha",
                "target",
                "status",
                "recorded_at",
                "details",
            }, f"ledger event keys differ on line {line_number}"
            event = SyncEvent(**raw)
            _validate_event(event)
            events.append(event)
        return tuple(events)

    def append(self, event: SyncEvent) -> None:
        """Append and fsync one event without rewriting prior lines."""
        _validate_event(event)
        self.path.parent.mkdir(parents=True, exist_ok=True)
        encoded = json.dumps(asdict(event), sort_keys=True, separators=(",", ":"))
        with self.path.open("a") as stream:
            stream.write(encoded + "\n")
            stream.flush()
            os.fsync(stream.fileno())


def _validate_event(event: SyncEvent) -> None:
    assert REQUEST_ID_RE.fullmatch(event.request_id)
    assert SHA_RE.fullmatch(event.source_sha)
    assert event.target in {"request", "website", "document"}
    assert event.status in {"registered", "succeeded", "failed"}
    assert ISO_UTC_RE.fullmatch(event.recorded_at)
    assert all(
        isinstance(key, str) and isinstance(value, str)
        for key, value in event.details.items()
    )
    if event.target == "request":
        assert event.status == "registered"
        assert set(event.details) == {"requested_at", "bundle_sha256"}
    elif event.status == "succeeded" and event.target == "website":
        assert set(event.details) == {"commit_sha", "preview_url", "build_status"}
        assert SHA_RE.fullmatch(event.details["commit_sha"])
        assert event.details["build_status"] == "success"
    elif event.status == "succeeded" and event.target == "document":
        assert set(event.details) == {
            "document_url",
            "revision_id",
            "verification_status",
        }
        assert event.details["verification_status"] == "success"
    else:
        assert set(event.details) == {"error"}


def _event_time() -> str:
    return datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")


def _events_for_request(ledger: SyncLedger, request_id: str) -> tuple[SyncEvent, ...]:
    return tuple(event for event in ledger.read() if event.request_id == request_id)


def _register_request(ledger: SyncLedger, request: SyncRequest) -> None:
    events = _events_for_request(ledger, request.request_id)
    registrations = [event for event in events if event.target == "request"]
    if registrations:
        identities = {
            (
                event.source_sha,
                event.details["requested_at"],
                event.details["bundle_sha256"],
            )
            for event in registrations
        }
        assert identities == {
            (request.source_sha, request.requested_at, request.bundle_sha256)
        }, f"request ID conflict: {request.request_id}"
        return
    ledger.append(
        SyncEvent(
            request_id=request.request_id,
            source_sha=request.source_sha,
            target="request",
            status="registered",
            recorded_at=_event_time(),
            details={
                "requested_at": request.requested_at,
                "bundle_sha256": request.bundle_sha256,
            },
        )
    )


def _successful_receipts(
    events: tuple[SyncEvent, ...],
) -> tuple[WebsiteReceipt | None, DocumentReceipt | None]:
    website_receipts = {
        WebsiteReceipt(
            commit_sha=event.details["commit_sha"],
            preview_url=event.details["preview_url"],
            build_status="success",
        )
        for event in events
        if event.target == "website" and event.status == "succeeded"
    }
    document_receipts = {
        DocumentReceipt(
            document_url=event.details["document_url"],
            revision_id=event.details["revision_id"],
            verification_status="success",
        )
        for event in events
        if event.target == "document" and event.status == "succeeded"
    }
    assert len(website_receipts) <= 1, "website receipt changed for one request"
    assert len(document_receipts) <= 1, "immutable document changed for one request"
    return (
        next(iter(website_receipts), None),
        next(iter(document_receipts), None),
    )


def _append_failure(
    ledger: SyncLedger,
    request: SyncRequest,
    target: Literal["website", "document"],
    error: Exception,
) -> str:
    message = f"{type(error).__name__}: {error}"
    ledger.append(
        SyncEvent(
            request_id=request.request_id,
            source_sha=request.source_sha,
            target=target,
            status="failed",
            recorded_at=_event_time(),
            details={"error": message},
        )
    )
    return message


def coordinate_sync(
    bundle: ReviewBundle,
    request: SyncRequest,
    *,
    ledger: SyncLedger,
    website: WebsitePublisher,
    document: DocumentPublisher,
) -> SyncReport:
    """Attempt only missing effects and retain enough state for safe retries."""
    assert bundle.manifest_path.is_file()
    assert _sha256(bundle.manifest_path) == request.bundle_sha256
    _register_request(ledger, request)
    events = _events_for_request(ledger, request.request_id)
    assert all(event.source_sha == request.source_sha for event in events)
    website_receipt, document_receipt = _successful_receipts(events)
    website_error: str | None = None
    document_error: str | None = None

    if website_receipt is None:
        try:
            website_receipt = website.publish(bundle, request)
            assert website_receipt.preview_url == WEBSITE_PREVIEW_URL
            ledger.append(
                SyncEvent(
                    request_id=request.request_id,
                    source_sha=request.source_sha,
                    target="website",
                    status="succeeded",
                    recorded_at=_event_time(),
                    details=asdict(website_receipt),
                )
            )
        except Exception as error:
            website_error = _append_failure(ledger, request, "website", error)

    if document_receipt is None:
        try:
            document_receipt = document.publish(bundle, request)
            ledger.append(
                SyncEvent(
                    request_id=request.request_id,
                    source_sha=request.source_sha,
                    target="document",
                    status="succeeded",
                    recorded_at=_event_time(),
                    details=asdict(document_receipt),
                )
            )
        except Exception as error:
            document_error = _append_failure(ledger, request, "document", error)

    return SyncReport(
        request=request,
        website=website_receipt,
        document=document_receipt,
        website_error=website_error,
        document_error=document_error,
    )


def request_from_bundle(bundle: ReviewBundle) -> SyncRequest:
    """Construct a sync request from a verified immutable bundle."""
    manifest = json.loads(bundle.manifest_path.read_text())
    return SyncRequest(
        request_id=manifest["request_id"],
        source_sha=manifest["source_sha"],
        requested_at=manifest["requested_at"],
        bundle_sha256=_sha256(bundle.manifest_path),
    )


def _parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--config", type=Path, default=default_config_path())
    subparsers = parser.add_subparsers(dest="command", required=True)
    prepare = subparsers.add_parser("prepare")
    prepare.add_argument("--source-sha", required=True)
    prepare.add_argument("--request-id", required=True)
    prepare.add_argument("--requested-at", required=True)
    prepare.add_argument("--destination", type=Path, required=True)
    return parser


def main() -> None:
    """Run the deterministic local preparation entry point."""
    args = _parser().parse_args()
    config = load_config(args.config)
    if args.command == "prepare":
        bundle = prepare_review_bundle(
            config,
            source_sha=args.source_sha,
            request_id=args.request_id,
            requested_at=args.requested_at,
            destination=args.destination,
        )
        request = request_from_bundle(bundle)
        print(
            json.dumps(
                {
                    "request_id": request.request_id,
                    "source_sha": request.source_sha,
                    "requested_at": request.requested_at,
                    "bundle_sha256": request.bundle_sha256,
                    "docx": str(bundle.docx_path),
                    "manifest": str(bundle.manifest_path),
                },
                indent=2,
                sort_keys=True,
            )
        )


if __name__ == "__main__":
    main()
